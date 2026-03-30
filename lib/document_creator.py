import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

DOCS_DIR = Path(r"C:\AI.Ass\data\workspace\documents")


class DocumentCreator:
    def __init__(self, config_path=r"C:\AI.Ass\config"):
        with open(Path(config_path) / "settings.json", encoding="utf-8") as f:
            self.config = json.load(f)
        DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _resolve_path(self, params: dict, ext: str) -> Path:
        filename = params.get("filename") or f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
        if not filename.endswith(f".{ext}"):
            filename += f".{ext}"
        return DOCS_DIR / filename

    def _success(self, doc_type: str, filepath: Path) -> dict:
        return {
            "status": "success",
            "doc_type": doc_type,
            "filepath": str(filepath),
            "message": f"{doc_type.upper()} saved to {filepath.name}"
        }

    def _error(self, reason: str, message: str) -> dict:
        return {"status": "error", "reason": reason, "message": message}

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def create_docx(self, content: str, params: dict = {}) -> dict:
        if not content or not isinstance(content, str):
            return self._error("invalid_content", "content must be a non-empty string")

        try:
            doc = Document()
            font_name = params.get("font", "Calibri")
            font_size = int(params.get("font_size", 12))
            heading_style = params.get("heading_style", "Heading 1")

            if params.get("title"):
                heading = doc.add_heading(params["title"], level=1)
                for run in heading.runs:
                    run.font.name = font_name

            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph()
                    continue

                if stripped.startswith("# "):
                    p = doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    p = doc.add_heading(stripped[3:], level=2)
                else:
                    p = doc.add_paragraph(stripped)
                    run = p.runs[0] if p.runs else p.add_run(stripped)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

            filepath = self._resolve_path(params, "docx")
            doc.save(filepath)
            return self._success("docx", filepath)
        except OSError as e:
            return self._error("file_write_error", str(e))
        except Exception as e:
            return self._error("docx_error", str(e))

    # ── XLSX ─────────────────────────────────────────────────────────────────

    def create_xlsx(self, data: dict, params: dict = {}) -> dict:
        if not isinstance(data, dict):
            return self._error("invalid_data", "data must be a dict with 'headers' and 'data' keys")

        headers = data.get("headers", [])
        rows = data.get("data", [])

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = params.get("sheet_name", "Sheet1")

            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(fill_type="solid", fgColor="2F5496")
            header_align = Alignment(horizontal="center", vertical="center")

            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                ws.column_dimensions[cell.column_letter].width = max(len(str(header)) + 4, 12)

            for row_idx, row in enumerate(rows, start=2):
                for col_idx, value in enumerate(row, start=1):
                    ws.cell(row=row_idx, column=col_idx, value=value)

            ws.freeze_panes = "A2"

            filepath = self._resolve_path(params, "xlsx")
            wb.save(filepath)
            return self._success("xlsx", filepath)
        except OSError as e:
            return self._error("file_write_error", str(e))
        except Exception as e:
            return self._error("xlsx_error", str(e))

    # ── PDF ──────────────────────────────────────────────────────────────────

    def create_pdf(self, content: str, params: dict = {}) -> dict:
        if not content or not isinstance(content, str):
            return self._error("invalid_content", "content must be a non-empty string")

        try:
            filepath = self._resolve_path(params, "pdf")
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                leftMargin=20 * mm,
                rightMargin=20 * mm,
                topMargin=20 * mm,
                bottomMargin=20 * mm
            )

            styles = getSampleStyleSheet()
            font_size = int(params.get("font_size", 12))

            body_style = ParagraphStyle(
                "body",
                parent=styles["Normal"],
                fontSize=font_size,
                leading=font_size * 1.4,
                spaceAfter=6
            )
            h1_style = ParagraphStyle(
                "h1",
                parent=styles["Heading1"],
                fontSize=font_size + 6,
                spaceBefore=12,
                spaceAfter=6
            )
            h2_style = ParagraphStyle(
                "h2",
                parent=styles["Heading2"],
                fontSize=font_size + 3,
                spaceBefore=10,
                spaceAfter=4
            )

            story = []

            if params.get("title"):
                story.append(Paragraph(params["title"], styles["Title"]))
                story.append(Spacer(1, 6 * mm))

            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 4 * mm))
                elif stripped.startswith("# "):
                    story.append(Paragraph(stripped[2:], h1_style))
                elif stripped.startswith("## "):
                    story.append(Paragraph(stripped[3:], h2_style))
                else:
                    story.append(Paragraph(stripped, body_style))

            doc.build(story)
            return self._success("pdf", filepath)
        except OSError as e:
            return self._error("file_write_error", str(e))
        except Exception as e:
            return self._error("pdf_error", str(e))
